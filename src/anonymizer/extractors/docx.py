from __future__ import annotations

from pathlib import Path
from typing import Iterator

from .base import BaseExtractor, UnsupportedFormatError
from ._markdown import to_markdown_table
from ..document_models import ExtractionResult


class DocxExtractor(BaseExtractor):
    """Extracts text from DOCX files.

    Paragraphs and tables are emitted **in document order**. Tables
    are rendered as GitHub-flavoured Markdown — preserves the column
    relationships through pseudonymization (markers replace PII inside
    cells without disturbing the table syntax).

    Each non-empty paragraph and each table becomes one DocumentBlock.
    Page numbers are not available from the python-docx model and are
    left as None.
    """

    supported_extensions = frozenset({".docx"})

    def extract(self, path: Path) -> ExtractionResult:
        try:
            from docx import Document
            from docx.oxml.ns import qn
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for DOCX extraction. "
                "Install it with: pip install 'python-docx>=1.1'"
            ) from exc

        try:
            doc = Document(str(path))
        except Exception as exc:
            raise UnsupportedFormatError(
                f"{path.name!r} could not be parsed as a DOCX file: {exc}"
            ) from exc

        raw: list[tuple[int | None, str]] = []
        for item in _iter_block_items(doc, qn, Paragraph, Table):
            if isinstance(item, Paragraph):
                raw.append((None, item.text))
            else:  # Table
                rows = [
                    [cell.text for cell in row.cells] for row in item.rows
                ]
                rendered = to_markdown_table(rows)
                if rendered:
                    raw.append((None, rendered))

        return self._build_result(raw)


def _iter_block_items(
    doc, qn, Paragraph, Table
) -> Iterator:
    """Yield paragraphs and tables from a document in body order.

    python-docx exposes ``doc.paragraphs`` and ``doc.tables`` as
    separate lists, losing the original order. To preserve it, we
    walk the body's children directly and check each XML tag.
    Recipe adapted from the python-docx documentation.
    """
    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    for child in body.iterchildren():
        if child.tag == p_tag:
            yield Paragraph(child, doc)
        elif child.tag == tbl_tag:
            yield Table(child, doc)
