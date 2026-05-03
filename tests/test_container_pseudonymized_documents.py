"""End-to-end tests for the pseudonymized-document import flow + the
``/validate-pseudonymized`` endpoint.

The most important guarantee here: importing an already-pseudonymized
document MUST NOT create new mapping entries. That's how the flow stays
distinct from the raw-document pipeline. Tests crash if that breaks.
"""
from __future__ import annotations

import json
import time
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from anonymizer_api.config import Settings
from anonymizer_api.main import create_app

POLICY_PATH = Path(__file__).parent.parent / "policies" / "default.yaml"


@pytest.fixture()
def api_settings(tmp_path: Path) -> Settings:
    return Settings(
        quarantine_dir=tmp_path / "quarantine",
        output_dir=tmp_path / "output",
        db_url=f"sqlite:///{tmp_path}/api.db",
        max_bytes=1 * 1024 * 1024,
        policy_path=POLICY_PATH,
        runtime_config_path=tmp_path / "runtime.json",
        use_mock_client=False,
        opf_use_mock_worker=True,
    )


@pytest.fixture()
def api_client(api_settings: Settings) -> TestClient:
    app = create_app(api_settings)
    with TestClient(app) as client:
        yield client


def _create_container(client: TestClient, name: str = "test") -> str:
    return client.post("/api/containers", json={"name": name}).json()[
        "container_id"
    ]


def _seed_with_raw_doc(client: TestClient, cid: str) -> None:
    """Populate the container's mapping with a raw-document upload.

    Sprint 5 flow: upload → wait for pending_review → approve →
    promotion runs. This helper drives the whole sequence so the
    tests can use the resulting markers.
    """
    body = (
        "Cliente: Joao Silva.\n"
        "Email: alice@example.com.\n"
    )
    files = {"file": ("doc.txt", body.encode("utf-8"), "text/plain")}
    r = client.post(
        f"/api/containers/{cid}/documents/raw", files=files
    )
    assert r.status_code == 201, r.text
    doc = r.json()
    # Wait for processing → pending_review
    deadline = time.monotonic() + 5.0
    while time.monotonic() < deadline:
        d = client.get(
            f"/api/containers/{cid}/documents/{doc['document_id']}"
        ).json()
        if d["status"] == "pending_review":
            break
        time.sleep(0.05)
    else:  # pragma: no cover
        raise AssertionError("seed: doc never reached pending_review")
    # Approve to trigger promotion
    r = client.post(f"/jobs/{doc['job_id']}/approve", json={})
    assert r.status_code == 200
    # Wait for ready
    deadline = time.monotonic() + 5.0
    while time.monotonic() < deadline:
        d = client.get(
            f"/api/containers/{cid}/documents/{doc['document_id']}"
        ).json()
        if d["status"] == "ready":
            return
        time.sleep(0.05)
    raise AssertionError("seed: doc never reached ready after approve")


def _upload_pseudonymized(
    client: TestClient,
    cid: str,
    *,
    body: str,
    filename: str = "edited.txt",
):
    files = {"file": (filename, body.encode("utf-8"), "text/plain")}
    return client.post(
        f"/api/containers/{cid}/documents/pseudonymized", files=files
    )


# ---------------------------------------------------------------------------
# Upload behaviour
# ---------------------------------------------------------------------------

class TestPseudonymizedUpload:
    def test_returns_pending_review_doc(
        self, api_client: TestClient
    ) -> None:
        """Pseudonymized uploads now go to ``pending_review`` — the
        operator approves via the dedicated review screen before the
        document becomes ``ready``."""
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)

        r = _upload_pseudonymized(
            api_client,
            cid,
            body="Resumo: [PESSOA_0001] e [EMAIL_0001].",
        )
        assert r.status_code == 201, r.text
        body = r.json()
        assert body["status"] == "pending_review"
        assert body["source_type"] == "already_pseudonymized_document"
        assert body["role"] == "edited_version"

    def test_does_not_create_mapping_entries(
        self, api_client: TestClient
    ) -> None:
        """The key invariant of the pseudonymized flow."""
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        before = api_client.get(f"/api/containers/{cid}/mapping").json()

        # Upload a pseudonymized doc that references TWO MARKERS that
        # don't exist in the container — they must NOT be added.
        _upload_pseudonymized(
            api_client,
            cid,
            body="Discussao: [PESSOA_9999] confirmou em [DOC_5555].",
        )

        after = api_client.get(f"/api/containers/{cid}/mapping").json()
        assert len(after) == len(before), (
            "Pseudonymized upload must not create mapping entries — "
            f"{len(before)} → {len(after)}."
        )

    def test_reports_known_unknown_and_malformed(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)

        r = _upload_pseudonymized(
            api_client,
            cid,
            body=(
                "Sumario: [PESSOA_0001] visitou [DOC_9999]. "
                "Notas: [joao_silva] mencionado."
            ),
        )
        assert r.status_code == 201, r.text
        document_id = r.json()["document_id"]

        # The detail endpoint surfaces the document fields, but the
        # validation summary is persisted on the row's
        # ``validation_summary_json`` column. We can read it back via
        # the normal GET (it's not exposed yet) — for now, sanity-check
        # via the dedicated validate-pseudonymized endpoint with the
        # same text.
        text = (
            "Sumario: [PESSOA_0001] visitou [DOC_9999]. "
            "Notas: [joao_silva] mencionado."
        )
        v = api_client.post(
            f"/api/containers/{cid}/validate-pseudonymized",
            json={"processed_text": text},
        ).json()

        assert "[PESSOA_0001]" in v["known_markers"]
        assert "[DOC_9999]" in v["unknown_markers"]
        assert "[joao_silva]" in v["malformed_markers"]
        assert v["is_clean"] is False
        # And the doc row should carry the same payload as JSON.
        assert document_id  # only used for sanity; full inspection omitted

    def test_rejects_unsupported_extension(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        r = _upload_pseudonymized(
            api_client,
            cid,
            body="x",
            filename="bad.exe",
        )
        assert r.status_code == 400

    def test_404_for_unknown_container(self, api_client: TestClient) -> None:
        r = _upload_pseudonymized(
            api_client,
            "00000000-0000-0000-0000-000000000000",
            body="anything",
        )
        assert r.status_code == 404


# ---------------------------------------------------------------------------
# /validate-pseudonymized — stateless
# ---------------------------------------------------------------------------

class TestValidatePseudonymizedEndpoint:
    def test_clean_text_against_seeded_container(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)

        r = api_client.post(
            f"/api/containers/{cid}/validate-pseudonymized",
            json={"processed_text": "[PESSOA_0001] reuniu [EMAIL_0001]."},
        )
        assert r.status_code == 200
        body = r.json()
        assert body["is_clean"] is True
        assert body["unknown_markers"] == []
        assert body["malformed_markers"] == []
        assert "[PESSOA_0001]" in body["known_markers"]

    def test_unknown_marker_is_reported(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)

        r = api_client.post(
            f"/api/containers/{cid}/validate-pseudonymized",
            json={"processed_text": "[PESSOA_9999] e [PESSOA_0001]"},
        )
        body = r.json()
        assert "[PESSOA_9999]" in body["unknown_markers"]
        assert body["is_clean"] is False

    def test_does_not_create_entries(self, api_client: TestClient) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        before = api_client.get(f"/api/containers/{cid}/mapping").json()

        api_client.post(
            f"/api/containers/{cid}/validate-pseudonymized",
            json={"processed_text": "[NEW_TYPE_0001] [PESSOA_0001]"},
        )

        after = api_client.get(f"/api/containers/{cid}/mapping").json()
        assert len(after) == len(before)

    def test_404_for_unknown_container(self, api_client: TestClient) -> None:
        r = api_client.post(
            "/api/containers/missing/validate-pseudonymized",
            json={"processed_text": "[PESSOA_0001]"},
        )
        assert r.status_code == 404


# ---------------------------------------------------------------------------
# Download regression — pseudonymised import must be downloadable as text
# even when the original upload was a binary format (.docx, .pdf).
# ---------------------------------------------------------------------------

class TestPseudonymizedDownload:
    @staticmethod
    def _approve(client: TestClient, cid: str, did: str) -> None:
        r = client.post(
            f"/api/containers/{cid}/documents/{did}/approve-pseudonymized"
        )
        assert r.status_code == 200, r.text

    def test_docx_import_downloads_as_utf8_text(
        self, api_client: TestClient
    ) -> None:
        """Regression: previously ``pseudonymized_path`` pointed at the
        binary upload, so ``GET /documents/{id}/download`` crashed with
        UnicodeDecodeError when the file was a .docx. Now the extracted
        text is saved as a separate ``.pseudonymized.txt`` artefact."""
        from io import BytesIO

        from docx import Document

        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)

        # Build a docx in-memory containing an already-pseudonymised
        # text. This is what a reviewer might paste back from an LLM.
        doc = Document()
        doc.add_paragraph("Sumário externo: [PESSOA_0001] revisou.")
        doc.add_paragraph("Contato: [EMAIL_0001].")
        buf = BytesIO()
        doc.save(buf)

        files = {
            "file": (
                "edited.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        r = api_client.post(
            f"/api/containers/{cid}/documents/pseudonymized", files=files
        )
        assert r.status_code == 201, r.text
        document_id = r.json()["document_id"]
        self._approve(api_client, cid, document_id)

        # Now download — must succeed and serve UTF-8 text, not raw docx.
        r = api_client.get(
            f"/api/containers/{cid}/documents/{document_id}/download"
        )
        assert r.status_code == 200, r.text
        assert r.headers["content-type"].startswith("text/plain")
        body = r.text  # would raise if response weren't valid UTF-8
        assert "[PESSOA_0001]" in body
        assert "[EMAIL_0001]" in body

    def test_txt_import_downloads_with_same_content(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        body = "Sumário: [PESSOA_0001] e [EMAIL_0001]."
        r = _upload_pseudonymized(api_client, cid, body=body)
        document_id = r.json()["document_id"]
        self._approve(api_client, cid, document_id)
        r = api_client.get(
            f"/api/containers/{cid}/documents/{document_id}/download"
        )
        assert r.status_code == 200
        assert r.text == body


# ---------------------------------------------------------------------------
# Pseudonymized review screen — full lifecycle
# ---------------------------------------------------------------------------

class TestPseudonymizedReviewLifecycle:
    def test_review_payload_includes_validation_and_residual(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        # Upload text mixing known marker + unknown marker + residual
        # (a fresh name not anonymised by whoever produced this text).
        r = _upload_pseudonymized(
            api_client,
            cid,
            body=(
                "Resumo: [PESSOA_0001] confirmou.\n"
                "Outra parte: [PESSOA_9999] tambem.\n"
                "Comentario: Roberto Lima opinou.\n"
            ),
        )
        document_id = r.json()["document_id"]

        r = api_client.get(
            f"/api/containers/{cid}/documents/{document_id}/review-pseudonymized"
        )
        assert r.status_code == 200, r.text
        body = r.json()
        assert body["document_id"] == document_id
        assert body["status"] == "pending_review"
        v = body["validation"]
        assert "[PESSOA_0001]" in v["known_markers"]
        assert "[PESSOA_9999]" in v["unknown_markers"]
        # Residual: the new name escaped pseudonymisation.
        residual_fragments = [r["fragment"] for r in body["residual_pii"]]
        assert any("Roberto" in f for f in residual_fragments), residual_fragments

    def test_residual_excludes_marker_internals(
        self, api_client: TestClient
    ) -> None:
        """The detector might lock onto ``PESSOA`` inside ``[PESSOA_0001]``
        but those overlaps must be filtered — they're part of the
        marker, not a leak."""
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        r = _upload_pseudonymized(
            api_client,
            cid,
            # Markers only, no residual content. PII detector might
            # still trigger on the bracketed labels — must be filtered.
            body="Texto: [PESSOA_0001] e [EMAIL_0001].",
        )
        document_id = r.json()["document_id"]
        r = api_client.get(
            f"/api/containers/{cid}/documents/{document_id}/review-pseudonymized"
        )
        body = r.json()
        # Every residual span must NOT overlap a marker.
        for span in body["residual_pii"]:
            text = body["text"]
            inside = text[span["start"] : span["end"]]
            assert "[" not in inside and "]" not in inside, (
                f"Residual span {span!r} overlaps a marker"
            )

    def test_approve_transitions_to_ready(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        r = _upload_pseudonymized(
            api_client,
            cid,
            body="OK: [PESSOA_0001].",
        )
        document_id = r.json()["document_id"]
        r = api_client.post(
            f"/api/containers/{cid}/documents/{document_id}/approve-pseudonymized"
        )
        assert r.status_code == 200
        assert r.json()["status"] == "ready"
        # After approval, download works.
        r = api_client.get(
            f"/api/containers/{cid}/documents/{document_id}/download"
        )
        assert r.status_code == 200

    def test_reject_transitions_to_rejected(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        r = _upload_pseudonymized(
            api_client,
            cid,
            body="Conteudo qualquer: [PESSOA_9999].",
        )
        document_id = r.json()["document_id"]
        r = api_client.post(
            f"/api/containers/{cid}/documents/{document_id}/reject-pseudonymized"
        )
        assert r.status_code == 200
        assert r.json()["status"] == "rejected"
        # Rejected docs cannot be downloaded.
        r = api_client.get(
            f"/api/containers/{cid}/documents/{document_id}/download"
        )
        assert r.status_code == 409

    def test_approve_rejects_doc_not_in_pending_review(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        r = _upload_pseudonymized(
            api_client,
            cid,
            body="OK: [PESSOA_0001].",
        )
        document_id = r.json()["document_id"]
        # First approval works
        r = api_client.post(
            f"/api/containers/{cid}/documents/{document_id}/approve-pseudonymized"
        )
        assert r.status_code == 200
        # Second time → 409 (status=ready)
        r = api_client.post(
            f"/api/containers/{cid}/documents/{document_id}/approve-pseudonymized"
        )
        assert r.status_code == 409


class TestPseudonymizedManualRedaction:
    def test_manual_redaction_creates_marker_and_replaces_text(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        r = _upload_pseudonymized(
            api_client,
            cid,
            body=(
                "Resumo: [PESSOA_0001] reuniu.\n"
                "Tambem participou: Roberto Lima.\n"
            ),
        )
        document_id = r.json()["document_id"]

        before_mapping = api_client.get(
            f"/api/containers/{cid}/mapping"
        ).json()
        before_persons = [
            m for m in before_mapping if m["entity_type"] == "private_person"
        ]

        r = api_client.post(
            f"/api/containers/{cid}/documents/{document_id}/manual-redaction-pseudonymized",
            json={
                "fragment": "Roberto Lima",
                "entity_type": "private_person",
            },
        )
        assert r.status_code == 200, r.text
        result = r.json()
        assert result["marker"].startswith("[PESSOA_")
        assert result["occurrences"] >= 1
        # Marker should be NEW (Roberto Lima wasn't in the seed).
        assert result["marker_created"] is True

        # The mapping table now has one extra entry.
        after_mapping = api_client.get(
            f"/api/containers/{cid}/mapping"
        ).json()
        after_persons = [
            m for m in after_mapping if m["entity_type"] == "private_person"
        ]
        assert len(after_persons) == len(before_persons) + 1

        # Re-fetch the review payload — the residual should be gone
        # and the new marker should be in the known list.
        r = api_client.get(
            f"/api/containers/{cid}/documents/{document_id}/review-pseudonymized"
        )
        body = r.json()
        assert "Roberto Lima" not in body["text"]
        assert result["marker"] in body["text"]
        assert result["marker"] in body["validation"]["known_markers"]

    def test_manual_redaction_reuses_existing_marker(
        self, api_client: TestClient
    ) -> None:
        """If the fragment was already mapped (e.g. from a previous
        raw-doc upload in the same container), the resolver reuses it
        instead of allocating a new one."""
        cid = _create_container(api_client)
        # Seed has 'Joao Silva' → [PESSOA_0001]
        _seed_with_raw_doc(api_client, cid)
        # Pseudonymized doc that accidentally has 'Joao Silva' in the
        # clear (the external editor missed it).
        r = _upload_pseudonymized(
            api_client,
            cid,
            body="Recap: Joao Silva atuou ainda.",
        )
        document_id = r.json()["document_id"]
        r = api_client.post(
            f"/api/containers/{cid}/documents/{document_id}/manual-redaction-pseudonymized",
            json={
                "fragment": "Joao Silva",
                "entity_type": "private_person",
            },
        )
        assert r.status_code == 200, r.text
        result = r.json()
        assert result["marker_created"] is False
        # And the marker is the one already in the mapping.
        assert result["marker"] == "[PESSOA_0001]"

    def test_manual_redaction_404_for_unknown_doc(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        r = api_client.post(
            f"/api/containers/{cid}/documents/missing/manual-redaction-pseudonymized",
            json={"fragment": "x", "entity_type": "private_person"},
        )
        assert r.status_code == 404

    def test_manual_redaction_409_when_already_approved(
        self, api_client: TestClient
    ) -> None:
        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)
        r = _upload_pseudonymized(
            api_client,
            cid,
            body="OK: [PESSOA_0001].",
        )
        document_id = r.json()["document_id"]
        api_client.post(
            f"/api/containers/{cid}/documents/{document_id}/approve-pseudonymized"
        )
        r = api_client.post(
            f"/api/containers/{cid}/documents/{document_id}/manual-redaction-pseudonymized",
            json={"fragment": "x", "entity_type": "private_person"},
        )
        assert r.status_code == 409


# ---------------------------------------------------------------------------
# Logging discipline — pseudonymized flow must not leak text or originals
# ---------------------------------------------------------------------------

class TestLoggingPrivacy:
    def test_pseudonymized_upload_no_pii_in_logs(
        self,
        api_client: TestClient,
        caplog: pytest.LogCaptureFixture,
    ) -> None:
        import logging

        cid = _create_container(api_client)
        _seed_with_raw_doc(api_client, cid)

        sentinel = "TOPSECRET_FRAGMENT_88"
        with caplog.at_level(logging.DEBUG):
            _upload_pseudonymized(
                api_client,
                cid,
                body=f"{sentinel} discutindo [PESSOA_0001].",
            )

        for record in caplog.records:
            assert sentinel not in record.getMessage(), (
                f"Pseudonymized upload leaked sentinel via "
                f"logger={record.name!r}"
            )
