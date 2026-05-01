"""Shared pytest fixtures for synthetic document files."""
from __future__ import annotations

from pathlib import Path

import pytest

_SYNTHETIC_TEXT = """\
SYNTHETIC SERVICE CONTRACT

This agreement is made between Jane Doe (born 15/06/1982) residing at
42 Placeholder Avenue, Test City, and Acme Corp.

Contact: jane.doe@synthetic-example.org | Phone: 555-0100-4321

Payment via account 4111-1111-1111-1111.

CONFIDENTIAL TOKEN: sk-abc123-fake-token-xyz789

Signature: Jane Doe
Date: 30/04/2026
"""


@pytest.fixture()
def synthetic_txt(tmp_path: Path) -> Path:
    p = tmp_path / "contract.txt"
    p.write_text(_SYNTHETIC_TEXT, encoding="utf-8")
    return p


@pytest.fixture()
def synthetic_md(tmp_path: Path) -> Path:
    p = tmp_path / "contract.md"
    p.write_text(
        "# Synthetic Contract\n\n" + _SYNTHETIC_TEXT,
        encoding="utf-8",
    )
    return p


@pytest.fixture()
def synthetic_docx(tmp_path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("SYNTHETIC SERVICE CONTRACT", level=1)
    doc.add_paragraph(
        "This agreement is made between Jane Doe (born 15/06/1982) "
        "residing at 42 Placeholder Avenue, Test City, and Acme Corp."
    )
    doc.add_paragraph(
        "Contact: jane.doe@synthetic-example.org | Phone: 555-0100-4321"
    )
    doc.add_paragraph("Payment via account 4111-1111-1111-1111.")
    doc.add_paragraph("CONFIDENTIAL TOKEN: sk-abc123-fake-token-xyz789")
    doc.add_paragraph("Signature: Jane Doe\nDate: 30/04/2026")

    path = tmp_path / "contract.docx"
    doc.save(str(path))
    return path


@pytest.fixture()
def synthetic_xlsx(tmp_path: Path) -> Path:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Contacts"
    ws.append(["Name", "Email", "Phone", "Account"])
    ws.append(["Jane Doe", "jane.doe@synthetic-example.org", "555-0100-4321", "4111-1111-1111-1111"])
    ws.append(["Bob Smith", "bob.smith@fake-corp.org", "555-0200-9999", "4222-2222-2222-2222"])

    path = tmp_path / "contacts.xlsx"
    wb.save(str(path))
    return path
