"""Unit tests for document extractors using synthetic in-memory files."""
from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from anonymizer.document_models import BLOCK_SEPARATOR
from anonymizer.extractors.base import UnsupportedFormatError
from anonymizer.extractors.docx import DocxExtractor
from anonymizer.extractors.pdf import PdfExtractor
from anonymizer.extractors.rtf import RtfExtractor
from anonymizer.extractors.txt import TxtExtractor
from anonymizer.extractors.xls import XlsExtractor
from anonymizer.extractors.xlsx import XlsxExtractor


# ---------------------------------------------------------------------------
# TXT / MD
# ---------------------------------------------------------------------------

class TestTxtExtractor:
    def test_extracts_paragraphs_as_blocks(self, synthetic_txt: Path) -> None:
        result = TxtExtractor().extract(synthetic_txt)
        assert len(result.blocks) >= 3
        assert all(b.text.strip() for b in result.blocks)

    def test_full_text_matches_joined_blocks(self, synthetic_txt: Path) -> None:
        result = TxtExtractor().extract(synthetic_txt)
        expected = BLOCK_SEPARATOR.join(b.text for b in result.blocks)
        assert result.full_text == expected

    def test_offsets_index_into_full_text(self, synthetic_txt: Path) -> None:
        result = TxtExtractor().extract(synthetic_txt)
        for block in result.blocks:
            assert result.full_text[block.start_offset : block.end_offset] == block.text

    def test_page_is_none_for_txt(self, synthetic_txt: Path) -> None:
        result = TxtExtractor().extract(synthetic_txt)
        assert all(b.page is None for b in result.blocks)

    def test_md_file_accepted(self, synthetic_md: Path) -> None:
        result = TxtExtractor().extract(synthetic_md)
        assert len(result.blocks) >= 1

    def test_binary_file_rejected(self, tmp_path: Path) -> None:
        binary = tmp_path / "bad.txt"
        binary.write_bytes(b"hello\x00world")
        with pytest.raises(UnsupportedFormatError, match="binary"):
            TxtExtractor().extract(binary)

    def test_empty_paragraphs_skipped(self, tmp_path: Path) -> None:
        p = tmp_path / "spaced.txt"
        p.write_text("Block A\n\n\n\nBlock B", encoding="utf-8")
        result = TxtExtractor().extract(p)
        assert len(result.blocks) == 2

    def test_block_ids_are_unique(self, synthetic_txt: Path) -> None:
        result = TxtExtractor().extract(synthetic_txt)
        ids = [b.block_id for b in result.blocks]
        assert len(ids) == len(set(ids))


# ---------------------------------------------------------------------------
# PDF (pypdf monkeypatched)
# ---------------------------------------------------------------------------

class TestPdfExtractor:
    def _make_mock_reader(self, page_texts: list[str]):
        pages = []
        for text in page_texts:
            page = MagicMock()
            page.extract_text.return_value = text
            pages.append(page)
        reader = MagicMock()
        reader.pages = pages
        return reader

    def test_one_block_per_page(self, tmp_path: Path) -> None:
        fake_path = tmp_path / "doc.pdf"
        fake_path.write_bytes(b"%PDF-1.4 fake")
        mock_reader = self._make_mock_reader(["Page one text.", "Page two text."])

        with patch("anonymizer.extractors.pdf.PdfReader", return_value=mock_reader):
            result = PdfExtractor().extract(fake_path)

        assert len(result.blocks) == 2
        assert result.blocks[0].page == 1
        assert result.blocks[1].page == 2

    def test_empty_pages_skipped(self, tmp_path: Path) -> None:
        fake_path = tmp_path / "doc.pdf"
        fake_path.write_bytes(b"%PDF-1.4 fake")
        mock_reader = self._make_mock_reader(["Text here.", "", "  "])

        with patch("anonymizer.extractors.pdf.PdfReader", return_value=mock_reader):
            result = PdfExtractor().extract(fake_path)

        assert len(result.blocks) == 1

    def test_offsets_index_into_full_text(self, tmp_path: Path) -> None:
        fake_path = tmp_path / "doc.pdf"
        fake_path.write_bytes(b"%PDF-1.4 fake")
        mock_reader = self._make_mock_reader(["Alpha beta.", "Gamma delta."])

        with patch("anonymizer.extractors.pdf.PdfReader", return_value=mock_reader):
            result = PdfExtractor().extract(fake_path)

        for block in result.blocks:
            assert result.full_text[block.start_offset : block.end_offset] == block.text

    def test_corrupt_pdf_raises_unsupported(self, tmp_path: Path) -> None:
        bad = tmp_path / "bad.pdf"
        bad.write_bytes(b"not a pdf at all")

        with patch(
            "anonymizer.extractors.pdf.PdfReader",
            side_effect=Exception("invalid PDF"),
        ):
            with pytest.raises(UnsupportedFormatError):
                PdfExtractor().extract(bad)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

class TestDocxExtractor:
    def test_extracts_paragraphs(self, synthetic_docx: Path) -> None:
        result = DocxExtractor().extract(synthetic_docx)
        assert len(result.blocks) >= 3

    def test_offsets_index_into_full_text(self, synthetic_docx: Path) -> None:
        result = DocxExtractor().extract(synthetic_docx)
        for block in result.blocks:
            assert result.full_text[block.start_offset : block.end_offset] == block.text

    def test_page_is_none(self, synthetic_docx: Path) -> None:
        result = DocxExtractor().extract(synthetic_docx)
        assert all(b.page is None for b in result.blocks)

    def test_corrupt_docx_raises_unsupported(self, tmp_path: Path) -> None:
        bad = tmp_path / "bad.docx"
        bad.write_bytes(b"not a docx")
        with pytest.raises(UnsupportedFormatError):
            DocxExtractor().extract(bad)

    def test_table_emitted_as_markdown_in_document_order(
        self, tmp_path: Path
    ) -> None:
        """A doc with paragraph → table → paragraph must produce three
        blocks in that order, with the middle block holding a Markdown
        table."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Antes da tabela.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Nome"
        table.cell(0, 1).text = "CPF"
        table.cell(1, 0).text = "Joao Silva"
        table.cell(1, 1).text = "111.444.777-35"
        doc.add_paragraph("Depois da tabela.")

        path = tmp_path / "with_table.docx"
        doc.save(str(path))

        result = DocxExtractor().extract(path)
        # Three blocks, body order preserved
        assert len(result.blocks) == 3
        assert result.blocks[0].text == "Antes da tabela."
        assert result.blocks[2].text == "Depois da tabela."
        # Middle block is the markdown table
        md = result.blocks[1].text
        assert "| Nome | CPF |" in md
        assert "| --- | --- |" in md
        assert "| Joao Silva | 111.444.777-35 |" in md


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------

class TestXlsxExtractor:
    def test_extracts_sheets_as_blocks(self, synthetic_xlsx: Path) -> None:
        result = XlsxExtractor().extract(synthetic_xlsx)
        assert len(result.blocks) >= 1

    def test_block_contains_cell_data(self, synthetic_xlsx: Path) -> None:
        result = XlsxExtractor().extract(synthetic_xlsx)
        combined = " ".join(b.text for b in result.blocks)
        assert "Jane Doe" in combined

    def test_offsets_index_into_full_text(self, synthetic_xlsx: Path) -> None:
        result = XlsxExtractor().extract(synthetic_xlsx)
        for block in result.blocks:
            assert result.full_text[block.start_offset : block.end_offset] == block.text

    def test_corrupt_xlsx_raises_unsupported(self, tmp_path: Path) -> None:
        bad = tmp_path / "bad.xlsx"
        bad.write_bytes(b"not xlsx")
        with pytest.raises(UnsupportedFormatError):
            XlsxExtractor().extract(bad)

    def test_sheet_rendered_as_markdown_table(
        self, synthetic_xlsx: Path
    ) -> None:
        result = XlsxExtractor().extract(synthetic_xlsx)
        assert result.blocks
        md = result.blocks[0].text
        # Header pipes from the conftest fixture (Name / Email / Phone / Account)
        assert "| Name |" in md
        assert "| --- |" in md
        # And the row data — Jane Doe sat in row 2
        assert "Jane Doe" in md

    def test_pipe_in_cell_is_escaped(self, tmp_path: Path) -> None:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "leakage"
        ws.append(["Categoria", "Anotação"])
        ws.append(["A | B", "valor com | pipe"])

        path = tmp_path / "pipe.xlsx"
        wb.save(str(path))

        result = XlsxExtractor().extract(path)
        md = result.blocks[0].text
        # The pipe in the cell content must be escaped so it doesn't
        # break the table.
        assert r"A \| B" in md
        assert r"valor com \| pipe" in md


# ---------------------------------------------------------------------------
# XLS (legacy Excel 97–2003)
# ---------------------------------------------------------------------------

class TestXlsExtractor:
    def test_extracts_sheets_as_blocks(self, synthetic_xls: Path) -> None:
        result = XlsExtractor().extract(synthetic_xls)
        assert len(result.blocks) >= 1

    def test_block_contains_cell_data(self, synthetic_xls: Path) -> None:
        result = XlsExtractor().extract(synthetic_xls)
        combined = " ".join(b.text for b in result.blocks)
        assert "Jane Doe" in combined
        assert "jane.doe@synthetic-example.org" in combined

    def test_offsets_index_into_full_text(self, synthetic_xls: Path) -> None:
        result = XlsExtractor().extract(synthetic_xls)
        for block in result.blocks:
            assert (
                result.full_text[block.start_offset : block.end_offset]
                == block.text
            )

    def test_sheet_rendered_as_markdown_table(
        self, synthetic_xls: Path
    ) -> None:
        result = XlsExtractor().extract(synthetic_xls)
        assert result.blocks
        md = result.blocks[0].text
        assert "| Name |" in md
        assert "| --- |" in md
        assert "Jane Doe" in md

    def test_page_index_is_sheet_number(self, synthetic_xls: Path) -> None:
        result = XlsExtractor().extract(synthetic_xls)
        # Single sheet → first block tagged as sheet 1.
        assert result.blocks[0].page == 1

    def test_corrupt_xls_raises_unsupported(self, tmp_path: Path) -> None:
        bad = tmp_path / "bad.xls"
        bad.write_bytes(b"not xls")
        with pytest.raises(UnsupportedFormatError):
            XlsExtractor().extract(bad)

    def test_pipe_in_cell_is_escaped(self, tmp_path: Path) -> None:
        import xlwt

        wb = xlwt.Workbook(encoding="utf-8")
        ws = wb.add_sheet("leakage")
        ws.write(0, 0, "Categoria")
        ws.write(0, 1, "Anotação")
        ws.write(1, 0, "A | B")
        ws.write(1, 1, "valor com | pipe")

        path = tmp_path / "pipe.xls"
        wb.save(str(path))

        result = XlsExtractor().extract(path)
        md = result.blocks[0].text
        assert r"A \| B" in md
        assert r"valor com \| pipe" in md


# ---------------------------------------------------------------------------
# RTF
# ---------------------------------------------------------------------------

# Minimal valid RTF fragment with two paragraphs and PT-BR content.
_SYNTHETIC_RTF = (
    r"{\rtf1\ansi\ansicpg1252\deff0"
    r"{\fonttbl{\f0 Times New Roman;}}"
    r"\f0\fs24 "
    r"Cliente: Joao Silva.\par "
    r"Email: alice@example.com.\par "
    r"}"
)


class TestRtfExtractor:
    def test_extracts_paragraphs(self, tmp_path: Path) -> None:
        path = tmp_path / "doc.rtf"
        path.write_text(_SYNTHETIC_RTF, encoding="utf-8")

        result = RtfExtractor().extract(path)
        # The synthetic doc has two visible paragraphs but striprtf may
        # collapse to one combined block — accept either.
        assert result.blocks, "should produce at least one block"
        full = " ".join(b.text for b in result.blocks)
        assert "Joao Silva" in full
        assert "alice@example.com" in full

    def test_offsets_index_into_full_text(self, tmp_path: Path) -> None:
        path = tmp_path / "doc.rtf"
        path.write_text(_SYNTHETIC_RTF, encoding="utf-8")

        result = RtfExtractor().extract(path)
        for block in result.blocks:
            assert (
                result.full_text[block.start_offset : block.end_offset]
                == block.text
            )

    def test_supported_extension(self) -> None:
        assert ".rtf" in RtfExtractor.supported_extensions

    def test_corrupt_rtf_does_not_crash(self, tmp_path: Path) -> None:
        """striprtf is permissive — it returns whatever text it can
        recover. The extractor must not raise on garbage input; at
        worst the result has no blocks."""
        path = tmp_path / "bad.rtf"
        path.write_text("not actually rtf", encoding="utf-8")
        result = RtfExtractor().extract(path)
        # Either empty or has the literal text; both are fine — the
        # detector layer handles whatever comes through.
        assert isinstance(result.full_text, str)


class TestPipelineRegistry:
    """Sanity: every supported extension maps to an extractor."""

    def test_all_expected_extensions_registered(self) -> None:
        from anonymizer.pipeline import ALLOWED_EXTENSIONS

        for ext in (
            ".txt",
            ".md",
            ".rtf",
            ".pdf",
            ".docx",
            ".xlsx",
            ".png",
            ".jpg",
            ".jpeg",
        ):
            assert ext in ALLOWED_EXTENSIONS, (
                f"{ext} should be a supported upload extension"
            )
